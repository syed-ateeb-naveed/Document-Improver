import time
import fitz
from docx import Document
from sumy.parsers.plaintext import PlaintextParser
from sumy.nlp.tokenizers import Tokenizer
from sumy.summarizers.lsa import LsaSummarizer
from tempfile import NamedTemporaryFile

def create_temp_file(uploaded_file):
    with NamedTemporaryFile(delete=False) as temp_file:
        for chunk in uploaded_file.chunks():
            temp_file.write(chunk)
        return temp_file.name
    
def count_words_in_docx(file_path):
    doc = Document(create_temp_file(file_path))
    total_words = 0

    for paragraph in doc.paragraphs:
        total_words += len(paragraph.text.split())

    return total_words

def count_words_in_pdf(file_path):
    doc = fitz.open(create_temp_file(file_path))
    total_words = 0

    for page_num in range(doc.page_count):
        page = doc.load_page(page_num)
        total_words += len(page.get_text("text").split())

    return total_words

def calculate_sentences_count(file_path):
    
    if file_path.name.endswith(".docx"):
        return max(1, count_words_in_docx(file_path) // 100)
    elif file_path.name.endswith(".pdf"):
        return max(1, count_words_in_pdf(file_path) // 100)
    else:
        raise ValueError("Unsupported file format")

def summarize(text, language="english", sentences_count=5):
    parser = PlaintextParser.from_string(text, Tokenizer(language))
    summarizer = LsaSummarizer()
    summary = summarizer(parser.document, sentences_count)
    return ' '.join([str(sentence) for sentence in summary])

def read_text_from_docx(file_path):
        
    try:
        doc = Document(create_temp_file(file_path))
        text = ""

        for paragraph in doc.paragraphs:
            text += paragraph.text + '\n'

        return text.strip()
    
    except Exception as e:
        print(f"Error reading PDF: {e}")
        return None


def read_text_from_pdf(file_path):

    try:
        doc = fitz.open(create_temp_file(file_path))
        text = ""

        for page_num in range(doc.page_count):
            page = doc.load_page(page_num)
            text += page.get_text("text") + '\n'

        return text.strip()
    
    except Exception as e:
        print(f"Error reading PDF: {e}")
        return None

if __name__ == "__main__":
    start_time = time.time()
    doc = Document("Sample.docx")
    sentences = calculate_sentences_count("Sample.docx")
    text = read_text_from_docx("Sample.docx")
    summary = summarize(text, sentences_count=sentences)
    print(summary)
    end_time = time.time()

    elapsed_time = end_time - start_time
    print(f"Elapsed time: {elapsed_time} seconds")
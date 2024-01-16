from django.shortcuts import render, HttpResponse, redirect
from django.contrib.auth import authenticate, login, logout
from django.contrib.auth.decorators import login_required
from django.contrib import messages
from django.contrib.auth.forms import UserCreationForm
from .forms import RegisterUserFrom
from .summarizer import summarize, read_text_from_docx, read_text_from_pdf, calculate_sentences_count
from .spelling import correct_spelling
from django.http import HttpResponse
from docx import Document
import io

# Create your views here.








def login_user(request):

    if request.method == "POST":
        username = request.POST["username"]
        password = request.POST["password"]
        user = authenticate(request, username=username, password=password)
        if user is not None:
            login(request, user)
            
            return redirect('index')
        else:
            messages.success(request, ("There was an error logging in :(("))
            return redirect('login')

    return render(request, 'login.html', {})

def logout_user(request):
    logout(request)
    messages.success(request, ("You were logged out"))

    return redirect('login')

def register_user(request):
    if request.method == "POST":
        form = RegisterUserFrom(request.POST)
        if form.is_valid():
            form.save()
            username = form.cleaned_data['username']
            password = form.cleaned_data['password1']
            user = authenticate(username=username, password=password)
            login(request, user)
            messages.success(request, ("Registrations Successfull"))
            return redirect('index')
    else:
        form = RegisterUserFrom()

    return render(request, 'register.html', {'form' : form})

@login_required
def index(request):

    return render(request, 'index.html', {})

@login_required
def summarization(request):
    if request.method == 'POST' and request.FILES.get('uploaded_file'):
        uploaded_file = request.FILES['uploaded_file']

        if uploaded_file.name.endswith('.docx'):
            text_content = read_text_from_docx(uploaded_file)
        elif uploaded_file.name.endswith('.pdf'):
            text_content = read_text_from_pdf(uploaded_file)
        else:
            raise ValueError("Unsupported file format")
        
        sentences_count = calculate_sentences_count(uploaded_file)
        summary = summarize(text_content, language="english", sentences_count=sentences_count)
        
        return render(request, 'summarization_result.html', {'summary': summary})

    return render(request, 'summarization_form.html')


@login_required
def spelling(request):
    if request.method == 'POST' and request.FILES.get('uploaded_file'):
        uploaded_file = request.FILES['uploaded_file']

        if uploaded_file.name.endswith('.docx'):
            text_content = read_text_from_docx(uploaded_file)
        elif uploaded_file.name.endswith('.pdf'):
            text_content = read_text_from_pdf(uploaded_file)
        else:
            raise ValueError("Unsupported file format")
        
      
        spelling = correct_spelling(text_content)

        doc = Document()

            # Add text to the document
        doc.add_heading('Corrected Spelling Document', level=1)
        doc.add_paragraph(spelling)  # Uncomment this line if you want to include the corrected text

            # Save the document to an in-memory file
        buffer = io.BytesIO()
        doc.save(buffer)

            # Create a response and set the content type to MS Word
        response = HttpResponse(buffer.getvalue(), content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')

            # Set the Content-Disposition header to force the browser to download the file
        response['Content-Disposition'] = 'attachment; filename=corrected_document.docx'
        return response
        
        return render(request, 'word_correction_result.html', {'spelling': spelling, 'download_link': response})

    return render(request, 'word_correction.html')



@login_required
def create_and_download_file(request):
    if request.method == 'POST' and request.FILES.get('uploaded_file'):
        uploaded_file = request.FILES['uploaded_file']

        if uploaded_file.name.endswith('.docx'):
            text_content = read_text_from_docx(uploaded_file)
        elif uploaded_file.name.endswith('.pdf'):
            text_content = read_text_from_pdf(uploaded_file)
        else:
            raise ValueError("Unsupported file format")
        
      
        spelling = correct_spelling(text_content)
        doc = Document()

            # Add text to the document
        doc.add_heading('Corrected Spelling Document', level=1)
        doc.add_paragraph(spelling)  # Uncomment this line if you want to include the corrected text

            # Save the document to an in-memory file
        buffer = io.BytesIO()
        doc.save(buffer)

            # Create a response and set the content type to MS Word
        response = HttpResponse(buffer.getvalue(), content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')

            # Set the Content-Disposition header to force the browser to download the file
        response['Content-Disposition'] = 'attachment; filename=corrected_document.docx'

        return response
        
     

    return render(request, 'word_correction.html')
